"""
Multi-format document ingestion (non-PDF).

LAYER: 1 (data-layer)
IMPORTS FROM: _models, _chunking, pdf_processing

Provides ingestion for plain text, Markdown, HTML, DOCX, EPUB, and a
universal ``ingest_file`` dispatcher that auto-detects format.  Also
includes ``detect_format`` for MIME/extension-based format identification.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from monitor_data.tools.ingest_tools._chunking import (
    _chunk_section,
    chunk_text,
)
from monitor_data.tools.ingest_tools._models import (
    _BS4_AVAILABLE,
    _DOCX_AVAILABLE,
    _EPUB_AVAILABLE,
    DEFAULT_MAX_TOKENS,
    DEFAULT_OVERLAP_RATIO,
    IngestedChunk,
)
from monitor_data.tools.ingest_tools.pdf_processing import ingest_pdf

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Internal format extractors
# ---------------------------------------------------------------------------


def _strip_markdown(text: str) -> str:
    """Convert Markdown markup to plain text (best-effort, no external deps)."""
    # Fenced code blocks — keep the code content
    text = re.sub(r"```[^\n]*\n(.*?)```", r"\1", text, flags=re.DOTALL)
    text = re.sub(r"`([^`\n]+)`", r"\1", text)
    # ATX headings
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Bold / italic
    text = re.sub(r"\*{1,3}([^*\n]+)\*{1,3}", r"\1", text)
    text = re.sub(r"_{1,3}([^_\n]+)_{1,3}", r"\1", text)
    # Links and images — keep alt/label text
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    # Blockquotes
    text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)
    # Horizontal rules
    text = re.sub(r"^[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
    # Table rows — strip pipe separators, keep cell content
    text = re.sub(
        r"^\|(.+)\|$",
        lambda m: "  ".join(c.strip() for c in m.group(1).split("|") if c.strip()),
        text,
        flags=re.MULTILINE,
    )
    # Alignment rows (|---|---|)
    text = re.sub(r"^\|[\s:|-]+\|$", "", text, flags=re.MULTILINE)
    # Collapse excess blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _strip_html(html: str) -> str:
    """Extract plain text from HTML, using BeautifulSoup when available."""
    if _BS4_AVAILABLE:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "head", "meta", "link", "noscript"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
    else:
        # Fallback: naive regex stripping
        text = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"&nbsp;", " ", text)
        text = re.sub(r"&amp;", "&", text)
        text = re.sub(r"&lt;", "<", text)
        text = re.sub(r"&gt;", ">", text)
        text = re.sub(r"&quot;", '"', text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_docx_text(docx_bytes: bytes) -> str:
    """Extract plain text from a .docx Word document paragraph by paragraph."""
    if not _DOCX_AVAILABLE:
        raise ImportError("python-docx is required for DOCX ingestion. Install it with: pip install python-docx")
    import io

    import docx

    doc = docx.Document(io.BytesIO(docx_bytes))
    parts: list[str] = []
    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            parts.append(stripped)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "  ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)


def _extract_epub_text(epub_bytes: bytes) -> str:
    """Extract plain text from an .epub e-book, chapter by chapter."""
    if not _EPUB_AVAILABLE:
        raise ImportError(
            "ebooklib and beautifulsoup4 are required for EPUB ingestion. "
            "Install with: pip install ebooklib beautifulsoup4"
        )
    import io

    import ebooklib
    from ebooklib import epub

    book = epub.read_epub(io.BytesIO(epub_bytes))
    sections: list[str] = []
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            html_content = item.get_content().decode("utf-8", errors="ignore")
            plain = _strip_html(html_content)
            if plain:
                sections.append(plain)
    return "\n\n".join(sections)


# ---------------------------------------------------------------------------
# Format detection
# ---------------------------------------------------------------------------

# Extension → canonical format name
_EXT_FORMAT_MAP: dict[str, str] = {
    ".pdf": "pdf",
    ".txt": "text",
    ".text": "text",
    ".rst": "text",  # reStructuredText — treat as plain text
    ".csv": "text",
    ".tsv": "text",
    ".log": "text",
    ".md": "markdown",
    ".markdown": "markdown",
    ".mdown": "markdown",
    ".mkd": "markdown",
    ".htm": "html",
    ".html": "html",
    ".xhtml": "html",
    ".docx": "docx",
    ".epub": "epub",
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".gif": "image",
    ".webp": "image",
    ".bmp": "image",
    ".tiff": "image",
    ".tif": "image",
    ".svg": "image",
}


def detect_format(filename: str, content_type: str | None = None) -> str:
    """
    Detect document format from filename extension or MIME type.

    Args:
        filename:     Original filename (e.g. "rulebook.pdf", "notes.md").
        content_type: Optional MIME type string for extra confidence
                      (e.g. "application/pdf", "text/html").

    Returns:
        One of: 'pdf', 'text', 'markdown', 'html', 'docx', 'epub', 'image', 'unknown'
    """
    ext = Path(filename).suffix.lower()
    if ext in _EXT_FORMAT_MAP:
        return _EXT_FORMAT_MAP[ext]

    # Fall back to MIME type content-type header
    if content_type:
        ct = content_type.lower().split(";")[0].strip()
        if "pdf" in ct:
            return "pdf"
        if ct == "text/html" or "html" in ct:
            return "html"
        if ct.startswith("text/"):
            return "text"
        if ct.startswith("image/"):
            return "image"
        if "wordprocessingml" in ct or "msword" in ct or "docx" in ct:
            return "docx"
        if "epub" in ct:
            return "epub"

    return "unknown"


# ---------------------------------------------------------------------------
# Multi-format public ingestion API
# ---------------------------------------------------------------------------


def _parse_markdown_sections(md_text: str) -> list[dict]:
    """
    Parse Markdown into heading-grouped sections using ATX headings (# to ######).

    Returns sections in the same format as ``_group_blocks_by_section()``:
        [{"headings": List[str], "body": [{"text": str, "page_number": int}]}]

    Fenced code blocks are protected — ``#`` inside code is not treated as a
    heading.  Body text is stripped of Markdown formatting while headings are
    preserved as structural metadata for ``section_path`` and
    ``semantic_category`` tagging.

    Returns an empty list when no ATX headings are found (caller should fall
    back to plain-text chunking).
    """
    lines = md_text.split("\n")
    heading_stack: list[str | None] = [None] * 6  # h1-h6
    sections: list[dict] = []
    current_body: list[dict] = []
    current_headings: list[str] = []
    body_lines: list[str] = []
    in_code_block = False
    found_any_heading = False

    def _flush_body() -> None:
        nonlocal body_lines
        if not body_lines:
            return
        raw = "\n".join(body_lines).strip()
        if raw:
            plain = _strip_markdown(raw)
            if plain.strip():
                current_body.append({"text": plain, "page_number": 0})
        body_lines = []

    def _flush_section() -> None:
        if current_body:
            sections.append({"headings": list(current_headings), "body": list(current_body)})
            current_body.clear()

    for line in lines:
        stripped = line.strip()

        # Track fenced code blocks — # inside code is code, not a heading
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            body_lines.append(line)
            continue

        if in_code_block:
            body_lines.append(line)
            continue

        # ATX heading: 1-6 hashes followed by a space then text
        heading_match = re.match(r"^(#{1,6})\s+(.+?)(?:\s+#*\s*)?$", line)
        if heading_match:
            found_any_heading = True
            _flush_body()
            _flush_section()

            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            # Strip inline formatting from heading text
            heading_text = re.sub(r"\*{1,3}([^*\n]+)\*{1,3}", r"\1", heading_text)
            heading_text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", heading_text)
            heading_text = re.sub(r"`([^`]+)`", r"\1", heading_text)

            idx = level - 1
            heading_stack[idx] = heading_text
            for i in range(idx + 1, len(heading_stack)):
                heading_stack[i] = None
            current_headings = [h for h in heading_stack if h]
        else:
            body_lines.append(line)

    _flush_body()
    _flush_section()

    return sections if found_any_heading else []


def ingest_text(
    text: str,
    source_name: str,
    *,
    max_tokens: int = DEFAULT_MAX_TOKENS,
    overlap_ratio: float = DEFAULT_OVERLAP_RATIO,
    metadata: dict | None = None,
) -> list[IngestedChunk]:
    """
    Chunk plain text (non-PDF) into embedding-ready pieces.

    Convenience wrapper around chunk_text() for use with raw text inputs
    (e.g. session notes, wiki articles, GM handouts in markdown).
    """
    return chunk_text(
        text,
        source_name,
        max_tokens=max_tokens,
        overlap_ratio=overlap_ratio,
        metadata=metadata,
    )


def ingest_markdown(
    md_text: str,
    source_name: str,
    *,
    max_tokens: int = DEFAULT_MAX_TOKENS,
    overlap_ratio: float = DEFAULT_OVERLAP_RATIO,
    metadata: dict | None = None,
) -> list[IngestedChunk]:
    """
    Parse Markdown headings into structured sections, then chunk with full
    ``section_path`` / ``semantic_category`` metadata (same quality as PDF
    ingestion).

    Falls back to plain-text chunking if no ATX headings are found.

    Args:
        md_text:      Raw Markdown content (e.g. from a .md file or wiki page).
        source_name:  Document identifier for provenance.
        max_tokens:   Maximum tokens per chunk (default 512).
        overlap_ratio: Token overlap fraction (default 10%).
        metadata:     Extra Qdrant payload fields.

    Returns:
        List of IngestedChunk objects ready for embedding.
    """
    sections = _parse_markdown_sections(md_text)
    if not sections:
        plain = _strip_markdown(md_text)
        return chunk_text(
            plain,
            source_name,
            max_tokens=max_tokens,
            overlap_ratio=overlap_ratio,
            metadata=metadata,
        )

    overlap_tokens = max(1, int(max_tokens * overlap_ratio))
    base_metadata = dict(metadata or {})
    all_chunks: list[IngestedChunk] = []

    for section in sections:
        section_chunks = _chunk_section(
            section,
            source_name,
            len(all_chunks),
            max_tokens,
            overlap_tokens,
            base_metadata,
        )
        all_chunks.extend(section_chunks)

    return all_chunks


def ingest_html(
    html_content: str,
    source_name: str,
    *,
    max_tokens: int = DEFAULT_MAX_TOKENS,
    overlap_ratio: float = DEFAULT_OVERLAP_RATIO,
    metadata: dict | None = None,
) -> list[IngestedChunk]:
    """
    Strip HTML tags and chunk content into embedding-ready pieces.

    Uses BeautifulSoup if installed (recommended), falls back to regex
    stripping otherwise.  Script/style blocks are removed before extraction.

    Args:
        html_content: Raw HTML string.
        source_name:  Document identifier for provenance.
        max_tokens:   Maximum tokens per chunk (default 512).
        overlap_ratio: Token overlap fraction (default 10%).
        metadata:     Extra Qdrant payload fields.

    Returns:
        List of IngestedChunk objects ready for embedding.
    """
    plain = _strip_html(html_content)
    return chunk_text(
        plain,
        source_name,
        max_tokens=max_tokens,
        overlap_ratio=overlap_ratio,
        metadata=metadata,
    )


def ingest_docx(
    docx_bytes: bytes,
    source_name: str,
    *,
    max_tokens: int = DEFAULT_MAX_TOKENS,
    overlap_ratio: float = DEFAULT_OVERLAP_RATIO,
    metadata: dict | None = None,
) -> list[IngestedChunk]:
    """
    Extract text from a .docx Word document and chunk into embedding-ready pieces.

    Extracts paragraphs and table cell text in document order.
    Requires: pip install python-docx

    Args:
        docx_bytes:   Raw .docx file bytes.
        source_name:  Document identifier for provenance.
        max_tokens:   Maximum tokens per chunk (default 512).
        overlap_ratio: Token overlap fraction (default 10%).
        metadata:     Extra Qdrant payload fields.

    Returns:
        List of IngestedChunk objects ready for embedding.

    Raises:
        ImportError: If python-docx is not installed.
    """
    plain = _extract_docx_text(docx_bytes)
    return chunk_text(
        plain,
        source_name,
        max_tokens=max_tokens,
        overlap_ratio=overlap_ratio,
        metadata=metadata,
    )


def ingest_epub(
    epub_bytes: bytes,
    source_name: str,
    *,
    max_tokens: int = DEFAULT_MAX_TOKENS,
    overlap_ratio: float = DEFAULT_OVERLAP_RATIO,
    metadata: dict | None = None,
) -> list[IngestedChunk]:
    """
    Extract text from an .epub e-book and chunk into embedding-ready pieces.

    Processes each EPUB document item, strips HTML, and concatenates chapter
    content before chunking.
    Requires: pip install ebooklib beautifulsoup4

    Args:
        epub_bytes:   Raw .epub file bytes.
        source_name:  Document identifier for provenance.
        max_tokens:   Maximum tokens per chunk (default 512).
        overlap_ratio: Token overlap fraction (default 10%).
        metadata:     Extra Qdrant payload fields.

    Returns:
        List of IngestedChunk objects ready for embedding.

    Raises:
        ImportError: If ebooklib is not installed.
    """
    plain = _extract_epub_text(epub_bytes)
    return chunk_text(
        plain,
        source_name,
        max_tokens=max_tokens,
        overlap_ratio=overlap_ratio,
        metadata=metadata,
    )


def ingest_file(
    file_bytes: bytes,
    filename: str,
    source_name: str | None = None,
    *,
    content_type: str | None = None,
    max_tokens: int = DEFAULT_MAX_TOKENS,
    overlap_ratio: float = DEFAULT_OVERLAP_RATIO,
    metadata: dict | None = None,
) -> list[IngestedChunk]:
    """
    Auto-detect document format and ingest into embedding-ready chunks.

    This is the primary universal entry point.  It routes to the correct
    extractor based on the filename extension (with MIME type as a fallback).

    Supported formats (sync):
        pdf, txt/rst/csv, md/markdown, html/htm, docx, epub

    Async-only formats (route to Indexer agent instead):
        images (.png/.jpg/.jpeg/.gif/.webp/.bmp) → Indexer.index_image()
        URIs (http/https)                        → Indexer.index_uri()

    Args:
        file_bytes:   Raw file bytes.
        filename:     Original filename — used for format detection.
        source_name:  Logical document name (defaults to the filename stem).
        content_type: Optional MIME type for extra format confidence.
        max_tokens:   Maximum tokens per chunk (default 512).
        overlap_ratio: Token overlap fraction (default 10%).
        metadata:     Extra Qdrant payload fields.

    Returns:
        List of IngestedChunk objects ready for embedding.

    Raises:
        ValueError: If format is 'image' — use Indexer.index_image() instead.
        ValueError: If format is unknown and bytes cannot be decoded as UTF-8 text.
    """
    fmt = detect_format(filename, content_type)
    sname = source_name or Path(filename).stem
    kwargs: dict = {
        "max_tokens": max_tokens,
        "overlap_ratio": overlap_ratio,
        "metadata": metadata,
    }

    if fmt == "pdf":
        return ingest_pdf(file_bytes, sname, **kwargs)
    if fmt == "text":
        return ingest_text(file_bytes.decode("utf-8", errors="replace"), sname, **kwargs)
    if fmt == "markdown":
        return ingest_markdown(file_bytes.decode("utf-8", errors="replace"), sname, **kwargs)
    if fmt == "html":
        return ingest_html(file_bytes.decode("utf-8", errors="replace"), sname, **kwargs)
    if fmt == "docx":
        return ingest_docx(file_bytes, sname, **kwargs)
    if fmt == "epub":
        return ingest_epub(file_bytes, sname, **kwargs)
    if fmt == "image":
        raise ValueError(
            "Images require async ingestion via LLM vision. "
            "Use Indexer.index_image() or Indexer.index_file() instead of ingest_file()."
        )

    # Unknown format — last resort: attempt UTF-8 text decoding
    try:
        text = file_bytes.decode("utf-8", errors="strict")
        return ingest_text(text, sname, **kwargs)
    except UnicodeDecodeError:
        raise ValueError(
            f"Cannot ingest '{filename}': unrecognised binary format. "
            f"Supported sync formats: pdf, txt, md, html, docx, epub. "
            f"For images use Indexer.index_image()."
        )
